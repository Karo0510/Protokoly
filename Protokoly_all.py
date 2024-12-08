from docx import Document
import docx
from docx.shared import Cm, Pt
from JSON_to_dict import JSON_to_dict
from datetime import datetime

#do zrobienia frontend
#oraz protokół zwrotu
#exe do zrobienia

test = r'C:\Users\090396\Desktop\template\test_6.docx'


input = r'C:\Users\090396\Desktop\template\szablon_calosc_2.docx'
output = r'C:\Users\090396\Desktop\template\Protokol_Kamila_Wiluszynska_(090442).docx'
path_g = r'C:\Users\090396\Desktop\template\kwm_p.json'
path_r = r''
#path_r = r''

name_font = "Arial"
size_paragraph = 11
table_size_paragraph = 10

cell_etiqueteG = ['"Asset Name" "Model"', '"Serial"', '"Asset Tag"']
cell_etiqueteR = ['"Item"', '"Serial"', '" "']

data1 = {
        '[AboutComputerG]': 'szyfrowanie Bitlocker',
        '[OperatingSystemG]': 'Microsoft Windows 11 pro x64',
        '[DesktopSoftwareG]': 'Microsoft Office Standard  2013',
        '[AntivirusSoftwareG]': 'ESET Endpoint Antivirus 11',
        '[OtherSoftwareG]':'7zip, Navigator',
        '[EmployeeIT]': 'Konrad Urbański',
        '[OrganizationEmployee]': 'Anna Pińczykowska',
        '[Given]':'21.11.2024',

        '[AboutComputerR]': '-',
        '[OperatingSystemR]': '',
        '[DesktopSoftwareR]': '',
        '[AntivirusSoftwareR]': '',
        '[OtherSoftwareR]':'',
        '[Return]':'21.11.2024'
    }

date_max = "2024-12-31 11:59PM"
date_min = "2024-11-18 01:00AM"


def create_row(input: str, output: str, number_of_row: int, index: int):

    if number_of_row <= 0:
        return
    
    doc = Document(input)

    table = doc.tables[index]

    table.add_row()

    for i in range(0, table._column_count):
        if index == 0:
            table.cell(number_of_row, i).text = cell_etiqueteR[i]
        elif index == 3:
            table.cell(number_of_row, i).text = cell_etiqueteG[i]


    doc.save(output)
    
    return

def table_style(input: str, output: str, *args):
    doc = Document(input)

    for table in doc.tables:
        print(table)
        for row in table.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size= Pt(table_size_paragraph)
                        font.name = name_font
                        run.text = run.text.replace('"', '')
  

    doc.save(output)

def paragraph_style(input: str, output: str):
    doc = Document(input)

    format_flag = True

    for paragraph in doc.paragraphs:
        if 'w:br w:type="page"' in paragraph._element.xml:
                break
        else:
            for run in paragraph.runs:                      
                print(run.text)
                run.font.size = Pt(size_paragraph)
                run.font.name = name_font
                run.text = run.text.replace('"', '')

    doc.save(output)

    return
    
    
#def create_table(input: str, output: str, number_of_rows: int):

    #doc = Document(input)

    #table = doc.tables[0]


    #for i in range(0, number_of_rows):
        #row = table.add_row()
        

    #for i in range(1, number_of_rows+1):
        #for j in range(0, table._column_count):
            #table.cell(i, j).text = cell_etiquete[j]
       
    #doc.save(output)
    
    #return

def fill_document(input: str, output: str, data: dict, column: str):
    doc = Document(input)

    if data is None:
        return

    for paragraph in doc.paragraphs:
        for key, value in data.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, "")

                if key == column:
                    paragraph.add_run(value).bold = True
                else:
                    paragraph.add_run(value).bold = False


    for i, table in enumerate(doc.tables):
        for row in table.rows:
            for cell in row.cells:
                for key, value in data.items():
                    if key in cell.text:
                        if (isinstance(value, str) == False):
                            value = str(value)
                        cell.text = cell.text.replace(key, '')
                        if (i == 2) or (i == 5):
                            cell.paragraphs[0].add_run(value).bold = True
                        else:
                            cell.paragraphs[0].add_run(value).bold = False


                       
    doc.save(output)
            
    return

def check_data(input: str):
    dt = datetime.strptime(input, "%Y-%m-%d %I:%M%p").date()
    return dt


def create_new_JSON_from_history(input: list):

    new_dict = []
    dmax = check_data(date_max)
    dmin = check_data(date_min)


    for index, l in enumerate(input):
        del l['Admin']
        for key, value in l.items():
            if value == "checkin from":
                dt = check_data(l["Date"])
                if (dt > dmax) or (dt < dmin):
                    continue;
                else:
                    new_dict.append(l)
                

    print(len(new_dict))

    return new_dict


def create_file(input: str, output: str):

    doc = Document(input)
    doc.save(output)


if __name__ == '__main__':

    create_file(input, output)

   

    if (path_r != ""):
        q = JSON_to_dict(path_r)
        data_r = q.JSON_to_data()
        l = create_new_JSON_from_history(data_r)
        print(l)

        for i in range(0, len(l)):
            create_row(output, output, i+1, 0)
            fill_document(output, output, l[i], "Checked Out To")


    if (path_g != ""):
        p = JSON_to_dict(path_g, None, None)
        data = p.JSON_to_data()
        print(data)

        for i in range(0, len(data)):
            create_row(output, output, i+1, 3)
            fill_document(output, output, data[i], "Checked Out To")


    fill_document(output, output, data1, "Checked Out To")

    paragraph_style(output, output)
    table_style(output, output)

    #try:
        #doc = docx.Document(test)
    #except:
        #doc = docx.Document()
        #doc.save(test)
        #print("Previous file was corrupted or didn't exist - new file was created.")
        #print("aaa")

    #p1 = doc.add_paragraph('This is a ')
    #p1.add_run('MS WORD ').bold = True
    #p1.add_run('document ')
    #eg = p1.add_run('example')
    #eg.font.size = Pt(20)
    #eg.font.color.rgb = RGBColor(0,128,0)

    #doc.save(test)
